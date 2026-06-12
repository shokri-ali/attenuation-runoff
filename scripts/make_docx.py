"""
make_docx.py — assemble manuscript/draft_v1.docx for review.
Sources: manuscript/00..06 *.md (HTML comments stripped), figures from
outputs/figures/, Table 1 from catchment_characteristics.csv, Table 2 built
in-script, Table 3 from figures/table1_main_results.csv.
Figures/tables are inserted right after the paragraph of their FIRST mention.

Rich markup (in md prose and in the strings below):
  *x*  -> italic        ~x~ -> subscript        ^x^ -> superscript
  (asterisks inside ~..~ / ^..^ make the sub/superscript italic)
Equations are re-typeset from the EQ table (italic variables, true sub/sup),
keyed by their trailing "(n)" in the md.
"""
from pathlib import Path
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT=Path(__file__).resolve().parent.parent
MS=ROOT/"manuscript"; FIG=ROOT/"outputs"/"figures"; OUT=ROOT/"outputs"
DOCX=MS/"draft_v1.docx"

TITLE="Do runoff models need rainfall? Direct attenuation–runoff modelling with a commercial microwave link"
AUTHORS="Saeid Esmaeil Nia, Ali Shokri*"
AFFIL="School of Engineering, The University of Waikato, Hamilton, New Zealand"
CORR="*Corresponding author"
DRAFT_NOTE="Internal review draft — 12 June 2026 (figures and tables placed at first mention)"

EQ={
 1:"*A*(*t*) = max( *B*~6h~(*t*) − *P*~avg~(*t*), 0 )",
 2:"*w*(*t*) = *α*~w~ · *w*(*t* − 1) + (1 − *α*~w~) · *A*(*t*)",
 3:"*c*(*t*) = *c*~0~ + (1 − *c*~0~) · *w*(*t*) / ( *w*(*t*) + *c*~ref~ )",
 4:"*u*(*t*) = *c*(*t*) · *A*(*t*)",
 5:"*s*~q~(*t*) = *a*~q~ · *s*~q~(*t* − 1) + (1 − *a*~q~) · *u*(*t*),     *s*~s~(*t*) = *a*~s~ · *s*~s~(*t* − 1) + (1 − *a*~s~) · *u*(*t*)",
 6:"*Q*~A~(*t*) = min( *k* [ *f* · *s*~q~(*t*) + (1 − *f*) · *s*~s~(*t*) ], *Q*~cap~ )",
 7:"*E*(*d*) = 3.0 + 2.3 · cos( 2π (*d* − 20) / 365 )  mm d^−1^",
 8:"*Q*~R~(*t*) = min( *c* · GR4J( *p*~s~ · *P*(*t*), *E*(*t*); *x*~1~…*x*~4~ ), *Q*~cap~ )",
 9:"*Q*~RA~(*t*) = *w* · *Q*~R~(*t*) + (1 − *w*) · *Q*~A~(*t*)",
 10:"*F* = 0.5 · NSE(*Q*) + 0.5 · NSE(√*Q*)",
}

CAPTIONS={
 "F1":("Figure 1.","Study area in eastern Melbourne, Australia: the commercial microwave link (red, 22.7 GHz, 3.8 km), the nine flow stations (triangles) and the Melbourne Water rain gauges (squares) on satellite imagery. The inset locates Melbourne within Australia.","F1_study_area.jpg",5.8),
 "F2":("Figure 2.","From received power to model forcing. (a) 15-min received power and its 6 h moving-average baseline over the full record; grey bands mark link outages, orange and green shading the calibration and validation periods. (b) One-week zoom on the 4–9 November 2018 storm: dips of received power below the local baseline (left axis) become the attenuation forcing *A* (right axis, shaded).","F2_signal_extraction.png",6.0),
 "F3":("Figure 3.","Effect of the de-baselining window length (2–24 h) on attenuation-only runoff skill: validation KGE and NSE averaged over the responsive catchments. KGE peaks for 5–6 h windows; the 6 h window is used throughout this study.","F3_window_sweep.png",5.0),
 "F4":("Figure 4.","Daily validation skill (KGE, top; NSE, bottom) of the rainfall benchmark (GR4J, green), the direct attenuation model (blue) and the output-level fusion (red) for the nine catchments, August–December 2018.","F4_headline_comparison.png",6.0),
 "F5":("Figure 5.","(a) Hourly attenuation forcing over the full simulation period; grey bands mark link outages, orange and green shading the calibration and validation periods. (b–d) Observed and simulated daily flow at Gardiner, Rowville and Wantirna South for the full period; legends quote each model's calibration / validation NSE.","F5_hydrographs.png",6.2),
 "F6":("Figure 6.","Attenuation-model hourly validation KGE versus (a) baseflow index and (b) flashiness (log scale) for the nine catchments; asterisks mark the regulated stations.","F6_catchment_controls.png",6.0),
 "F7":("Figure 7.","Hourly performance of the attenuation-only model. (a) Validation KGE per catchment at hourly and daily resolution. (b) Scoresby, 6 November 2018 storm and (c) 11 August 2018 winter event: hourly observed flow (black), hourly attenuation-only simulation (blue) and the observed daily mean (dashed steps), showing the information removed by daily aggregation.","F7_hourly_attenuation.png",6.2),
 "F8":("Figure 8.","Hourly observed versus simulated flow (attenuation-only model) for all validation hours in all nine catchments on log–log axes; solid line 1:1, dashed lines 5:1 and 1:5. Grey points are link-outage hours, when the forcing is imposed as dry; flows below 0.005 m³ s⁻¹ are clipped to the axis floor.","F8_obs_sim_scatter_hourly.png",5.8),
}

INSERTS=[
 ("convective storm of 6 November 2018", ["F1","T1"]),
 ("enters the runoff model directly",    ["F2"]),
 ("applied unchanged at hourly and daily time steps", ["T2"]),
 ("zero forcing throughout this outage", ["F3"]),
 ("(NSE 0.52, KGE 0.73)",                ["T3","F4"]),
 ("visibly depressing the validation scores", ["F5"]),
 ("aggregation removed the timing differences", ["F6"]),
 ("rain-driven events occurred",         ["F7","F8"]),
]

doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)

TOK=re.compile(r"(\*[^*]+\*)|(~[^~]+~)|(\^[^^]+\^)")
def rich(p,text,size=None,bold=False):
    """Render markup text into runs on paragraph p."""
    pos=0
    for m in TOK.finditer(text):
        if m.start()>pos: _run(p,text[pos:m.start()],size,bold)
        tok=m.group(0); inner=tok[1:-1]
        ital = tok[0]=="*" or "*" in inner
        inner=inner.replace("*","")
        r=_run(p,inner,size,bold)
        r.italic=ital
        if tok[0]=="~": r.font.subscript=True
        if tok[0]=="^": r.font.superscript=True
        pos=m.end()
    if pos<len(text): _run(p,text[pos:],size,bold)
def _run(p,text,size,bold):
    r=p.add_run(text)
    if size: r.font.size=Pt(size)
    if bold: r.bold=True
    return r

def para(text,align=None,italic=False,bold=False,size=None,color=None):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.italic=italic; r.bold=bold
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor(*color)
    if align is not None: p.alignment=align
    return p

def caption(label,text):
    p=doc.add_paragraph()
    r=p.add_run(label+" "); r.bold=True; r.font.size=Pt(9)
    rich(p,text,size=9)

def add_figure(key):
    lab,cap,fn,w=CAPTIONS[key]
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG/fn),width=Inches(w))
    caption(lab,cap)

def add_equation(n):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rich(p,EQ[n])
    r=p.add_run("        ("+str(n)+")")
    r.font.size=Pt(11)

def style_table(t):
    t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(t.rows):
        for c in row.cells:
            for pp in c.paragraphs:
                for r in pp.runs:
                    r.font.size=Pt(8.5)
                    if i==0: r.bold=True

def cell_rich(cell,text):
    cell.text=""
    rich(cell.paragraphs[0],text,size=8.5)

def add_table1():
    caption("Table 1.","Characteristics of the nine study catchments over 1 November 2017 – 19 December 2018: distance from the flow station to the link midpoint, mean and maximum hourly discharge, flashiness (Qmax/Qmean), baseflow index (BFI), and the percentage of zero-flow hours.")
    d=pd.read_csv(OUT/"catchment_characteristics.csv")
    hdr=["Catchment","Distance to link (km)","Qmean (m³ s⁻¹)","Qmax (m³ s⁻¹)","Flashiness","BFI","Zero-flow hours (%)"]
    t=doc.add_table(rows=1,cols=len(hdr))
    for j,h in enumerate(hdr): t.rows[0].cells[j].text=h
    for _,r in d.iterrows():
        cells=t.add_row().cells
        vals=[r["catchment"],f"{r['dist_km']:.1f}",f"{r['Qmean']:.3f}",f"{r['Qmax']:.1f}",f"{r['flashiness']:.0f}",f"{r['BFI']:.2f}",f"{r['zeroflow_%']:.1f}"]
        for j,v in enumerate(vals): cells[j].text=str(v)
    style_table(t)

def add_table2():
    caption("Table 2.","Calibrated parameters and bounds. Transfer model parameters are dimensionless except *k*; the attenuation forcing is normalised by its calibration-period standard deviation, so *k* absorbs all amplitude scaling.")
    rows=[
     ("Attenuation transfer model","",""),
     ("*α*~w~","antecedent-wetness memory (Eq. 2)","0.5–0.999"),
     ("*c*~ref~","wetness half-saturation (Eq. 3)","0.01–10"),
     ("*c*~0~","immediate runoff-coefficient floor (Eq. 3)","0–1"),
     ("*a*~q~","quick-store recession (Eq. 5)","0–0.95"),
     ("*a*~s~","slow-store recession (Eq. 5)","0.90–0.99"),
     ("*f*","quick/slow partition (Eq. 6)","0–1"),
     ("*k*","output scale, m³ s⁻¹ per unit forcing (Eq. 6)","0–15"),
     ("GR4J rainfall benchmark","",""),
     ("*x*~1~","production store capacity (mm)","50–2000"),
     ("*x*~2~","groundwater exchange (mm d^−1^)","−5–5"),
     ("*x*~3~","routing store capacity (mm)","10–500"),
     ("*x*~4~","unit-hydrograph time base (d)","0.5–10"),
     ("*p*~s~","rainfall multiplier (Eq. 8)","0.5–5"),
     ("*c*","output scale, m³ s⁻¹ mm^−1^ (Eq. 8)","0–10"),
     ("Fusion","",""),
     ("*w*","weight on the rainfall prediction (Eq. 9)","0–1"),
    ]
    t=doc.add_table(rows=1,cols=3)
    for j,h in enumerate(["Parameter","Meaning","Bounds"]): t.rows[0].cells[j].text=h
    for a,b,c in rows:
        cells=t.add_row().cells
        cell_rich(cells[0],a); cell_rich(cells[1],b); cell_rich(cells[2],c)
    style_table(t)

def add_table3():
    caption("Table 3.","Daily results (best of five calibration seeds) for the rainfall benchmark (R), the direct attenuation model (A) and the output-level fusion (RA): calibration NSE, validation NSE and validation KGE per pathway, and the calibrated fusion weight *w* (*w* = 1 is fully rain-driven). All validation metrics are out of sample (August–December 2018).")
    d=pd.read_csv(FIG/"table1_main_results.csv")
    hdr=["Catchment","NSEcal R","NSE R","KGE R","NSEcal A","NSE A","KGE A","NSEcal RA","NSE RA","KGE RA","*w*"]
    t=doc.add_table(rows=1,cols=len(hdr))
    for j,h in enumerate(hdr): cell_rich(t.rows[0].cells[j],h)
    for _,r in d.iterrows():
        cells=t.add_row().cells
        vals=[r["catchment"]]+[f"{r[c]:.2f}" for c in ["NSEcal_R","NSE_R","KGE_R","NSEcal_A","NSE_A","KGE_A","NSEcal_RA","NSE_RA","KGE_RA"]]+[f"{r['fusion_w']:.2f}"]
        for j,v in enumerate(vals): cells[j].text=str(v)
    style_table(t)

TABLES={"T1":add_table1,"T2":add_table2,"T3":add_table3}

# ---------- title block ----------
para(TITLE,align=WD_ALIGN_PARAGRAPH.CENTER,bold=True,size=15)
para(AUTHORS,align=WD_ALIGN_PARAGRAPH.CENTER,size=11)
para(AFFIL,align=WD_ALIGN_PARAGRAPH.CENTER,italic=True,size=10)
para(CORR,align=WD_ALIGN_PARAGRAPH.CENTER,size=9)
para(DRAFT_NOTE,align=WD_ALIGN_PARAGRAPH.CENTER,italic=True,size=9,color=(150,30,30))
doc.add_paragraph()

# ---------- body ----------
eq_re=re.compile(r"\((\d+)\)\s*$")
done=set()
# HESS section order: body -> code/data availability etc. (07) -> references (06)
for f in ["00_abstract.md","01_introduction.md","02_methods.md","03_results.md","04_discussion.md","05_conclusions.md","07_backmatter.md","06_references.md"]:
    txt=(MS/f).read_text(encoding="utf-8")
    txt=re.sub(r"<!--.*?-->","",txt,flags=re.S)
    for block in [b.strip() for b in txt.split("\n\n")]:
        if not block or block=="---": continue
        if block.startswith("## "):
            doc.add_heading(block[3:].strip(),level=2); continue
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(),level=1); continue
        block_one=" ".join(l.strip() for l in block.splitlines())
        m=eq_re.search(block_one)
        if m and "=" in block_one and len(block_one)<240 and int(m.group(1)) in EQ:
            add_equation(int(m.group(1)))
        else:
            p=doc.add_paragraph()
            rich(p,block_one)
            p.paragraph_format.space_after=Pt(8)
        for anchor,items in INSERTS:
            if anchor in block_one and anchor not in done:
                done.add(anchor)
                for it in items:
                    if it.startswith("T"): TABLES[it]()
                    else: add_figure(it)
                    doc.add_paragraph()

missing=[a for a,_ in INSERTS if a not in done]
if missing: print("WARNING — anchors not found:",missing)
doc.save(DOCX)
print("saved:",DOCX)
